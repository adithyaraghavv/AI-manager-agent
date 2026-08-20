import io

from docx import Document

from app.core.text_extraction import extract_text


def test_extract_text_from_txt():
    content = "Contract value: $50,000\nScope: build a widget.".encode("utf-8")
    assert (
        extract_text(content, "txt")
        == "Contract value: $50,000\nScope: build a widget."
    )


def test_extract_text_from_txt_handles_bad_bytes_by_replacing_not_raising():
    content = b"\xff\xfe not valid utf-8"
    result = extract_text(content, "txt")
    assert result is not None


def test_extract_text_from_docx():
    document = Document()
    document.add_paragraph("Statement of Work")
    document.add_paragraph("Contract value: $75,000")
    buffer = io.BytesIO()
    document.save(buffer)

    result = extract_text(buffer.getvalue(), "docx")

    assert "Statement of Work" in result
    assert "Contract value: $75,000" in result


def test_extract_text_unsupported_extension_returns_none():
    assert extract_text(b"whatever bytes", "xlsx") is None
    assert extract_text(b"whatever bytes", "pptx") is None
    assert extract_text(b"whatever bytes", "doc") is None


def test_extract_text_from_garbage_pdf_bytes_returns_none_not_raises():
    # A corrupted/non-PDF file with a .pdf extension must fail gracefully —
    # callers treat "no text" as one outcome, not a crash.
    assert extract_text(b"this is not a real pdf file", "pdf") is None


def test_extract_text_from_empty_content_returns_none():
    assert extract_text(b"", "txt") is None
